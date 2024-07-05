import pytest
from src.docx_replace_package_alexgarden92160.app import generate_word, parse_form_data
from docx import Document
import io

def extract_text_from_doc(doc):
    text = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text.append(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text.append(run.text)
    return text

def test_generate_word():
    template_path = "templates/1.docx"
    form_data = {"name": "John Doe", "ville": "Versailles"}
    
    buffer = generate_word(template_path, form_data)
    
    buffer.seek(0)
    doc = Document(buffer)
    
    # Check if the placeholders were replaced correctly
    doc_text = extract_text_from_doc(doc)
    assert "John Doe" in doc_text
    assert "Versailles" in doc_text

def test_parse_form_data():
    form_data_str = "name=John Doe, ville=Versailles"
    expected_output = {"name": "John Doe", "ville": "Versailles"}
    
    parsed_data = parse_form_data(form_data_str)
    
    assert parsed_data == expected_output

@pytest.fixture
def client():
    from src.docx_replace_package_alexgarden92160.app import app
    app.config['TESTING'] = True
    with app.test_client() as client:
        yield client

def test_generate_endpoint(client):
    response = client.post('/generate', json={
        "template_path": "templates/1.docx",
        "form_data": {"name": "John Doe", "ville": "Versailles"}
    })
    
    assert response.status_code == 200
    assert response.headers["Content-Disposition"] == "attachment; filename=document.docx"

    # Check if the content is a valid Word document
    buffer = io.BytesIO(response.data)
    doc = Document(buffer)
    doc_text = extract_text_from_doc(doc)
    assert "John Doe" in doc_text
    assert "Versailles" in doc_text
