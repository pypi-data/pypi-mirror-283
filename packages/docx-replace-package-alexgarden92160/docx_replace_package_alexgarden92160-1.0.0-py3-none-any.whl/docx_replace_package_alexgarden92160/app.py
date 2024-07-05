from flask import Flask, request, send_file
from docx import Document
import io
import argparse

app = Flask(__name__)

def generate_word(template_path, form_data):
    # Load the template Word document
    doc = Document(template_path)
    # Replace placeholders with form data
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for field, value in form_data.items():
                if f'{{{{ {field} }}}}' in run.text:
                    run.text = run.text.replace(f'{{{{ {field} }}}}', str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for field, value in form_data.items():
                            if f'{{{{ {field} }}}}' in run.text:
                                run.text = run.text.replace(f'{{{{ {field} }}}}', str(value))

    # Save to a temporary file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/generate', methods=['POST'])
def generate():
    data = request.json
    template_path = data.get('template_path')
    form_data = data.get('form_data')

    buffer = generate_word(template_path, form_data)
    
    return send_file(buffer, as_attachment=True, download_name='document.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def parse_form_data(form_data_str):
    form_data = {}
    items = form_data_str.split(',')
    for item in items:
        key, value = item.split('=')
        form_data[key.strip()] = value.strip()
    return form_data

def cli():
    parser = argparse.ArgumentParser(description="Generate a Word document from a template.")
    parser.add_argument('--cli', action='store_true', help="Use the command line interface")
    parser.add_argument('template_path', type=str, help="Path to the template to use")
    parser.add_argument('form_data', type=str, help="String with the form data (e.g. 'name=John Doe, age=30')")
    parser.add_argument('output', type=str, help="Output Word file path")

    args = parser.parse_args()

    if args.cli:
        template_path = args.template_path
        form_data = parse_form_data(args.form_data)
        output = args.output

        buffer = generate_word(template_path, form_data)

        with open(output, 'wb') as f:
            f.write(buffer.getbuffer())
    else:
        app.run(host='0.0.0.0', port=5000)

if __name__ == '__main__':
    cli()
